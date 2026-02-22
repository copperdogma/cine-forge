from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.shared import Inches, Pt
from fpdf import FPDF


class ScreenplayPDF(FPDF):
    def __init__(self):
        super().__init__(orientation="P", unit="in", format="Letter")
        self.set_margins(left=1.5, top=1.0, right=1.0)
        self.set_auto_page_break(auto=True, margin=1.0)
        self.set_font("Courier", size=12)
        self.line_height = 1/6

    def header(self):
        # Page numbers in top right (starting from script page 2 = PDF page 3)
        if self.page_no() > 2:
            self.set_y(0.5)
            self.set_x(-1.5)
            self.set_font("Courier", "", 12)
            self.cell(0, 0, f"{self.page_no() - 1}.", align="R")
            self.set_y(1.0)

    def footer(self):
        pass

    def render_title_page(self, title_lines: list[str]):
        if not title_lines:
            return
        self.add_page()
        self.set_font("Courier", "", 12)
        self.set_y(3.5) # Start 1/3 down
        
        for i, line in enumerate(title_lines):
            if not line.strip():
                self.ln(self.line_height)
                continue
            if i == 0:
                self.set_font("Courier", "B", 12)
                self.multi_cell(
                    0, self.line_height, line.strip().upper(), 
                    align="C", new_x="LMARGIN", new_y="NEXT"
                )
                self.set_font("Courier", "", 12)
            else:
                self.multi_cell(
                    0, self.line_height, line.strip(), 
                    align="C", new_x="LMARGIN", new_y="NEXT"
                )

class ScreenplayRenderer:
    def sanitize(self, text: str | Any) -> str:
        if text is None:
            return ""
        s = str(text)
        replacements = {
            "\u2018": "'", "\u2019": "'",
            "\u201c": '"', "\u201d": '"',
            "\u2013": "-", "\u2014": "-",
            "\u2026": "...",
        }
        for k, v in replacements.items():
            s = s.replace(k, v)
        return s.encode("latin-1", "replace").decode("latin-1")

    def _split_pre_scene(self, text: str) -> tuple[list[str], list[str]]:
        lines = text.splitlines()
        title_lines = []
        teaser_lines = []
        script_started = False
        markers = ["TEASER", "ACT ONE", "FADE IN", "EXT.", "INT."]
        
        for line in lines:
            if not script_started:
                upper_line = line.strip().upper()
                if any(m in upper_line for m in markers):
                    script_started = True
                    teaser_lines.append(line)
                else:
                    title_lines.append(line)
            else:
                teaser_lines.append(line)
        return title_lines, teaser_lines

    def render_pdf(
        self, scenes: list[dict[str, Any]], output_path: str, 
        pre_scene_text: str = "", project_title: str = "Untitled"
    ):
        pdf = ScreenplayPDF()
        title_lines, teaser_lines = self._split_pre_scene(pre_scene_text)
        
        if title_lines:
            pdf.render_title_page(title_lines)
        else:
            pdf.render_title_page([project_title])
        
        pdf.add_page()
        if teaser_lines:
            pdf.set_font("Courier", "", 12)
            for line in teaser_lines:
                if not line.strip():
                    pdf.ln(pdf.line_height)
                    continue
                pdf.multi_cell(
                    0, pdf.line_height, self.sanitize(line), 
                    align="L", new_x="LMARGIN", new_y="NEXT"
                )
            pdf.ln(pdf.line_height)

        for scene in scenes:
            for elem in scene.get("elements", []):
                etype = elem.get("element_type", "action")
                content = self.sanitize(elem.get("content", ""))
                if etype == "scene_heading":
                    pdf.ln(pdf.line_height)
                    pdf.set_font("Courier", "B", 12)
                    pdf.set_x(1.5)
                    pdf.multi_cell(
                        0, pdf.line_height, content.upper(), 
                        align="L", new_x="LMARGIN", new_y="NEXT"
                    )
                    pdf.ln(pdf.line_height)
                elif etype == "character":
                    pdf.ln(pdf.line_height)
                    pdf.set_font("Courier", "", 12)
                    pdf.set_x(3.5)
                    pdf.multi_cell(
                        2.0, pdf.line_height, content.upper(), 
                        align="L", new_x="LMARGIN", new_y="NEXT"
                    )
                elif etype == "parenthetical":
                    pdf.set_font("Courier", "", 12)
                    pdf.set_x(3.0)
                    pdf.multi_cell(
                        2.0, pdf.line_height, f"({content})", 
                        align="L", new_x="LMARGIN", new_y="NEXT"
                    )
                elif etype == "dialogue":
                    pdf.set_font("Courier", "", 12)
                    pdf.set_x(2.5)
                    pdf.multi_cell(
                        3.5, pdf.line_height, content, 
                        align="L", new_x="LMARGIN", new_y="NEXT"
                    )
                    pdf.ln(pdf.line_height)
                elif etype == "transition":
                    pdf.ln(pdf.line_height)
                    pdf.set_font("Courier", "", 12)
                    pdf.set_x(5.5)
                    pdf.multi_cell(
                        2.0, pdf.line_height, content.upper(), 
                        align="L", new_x="LMARGIN", new_y="NEXT"
                    )
                    pdf.ln(pdf.line_height)
                else: # action
                    pdf.set_font("Courier", "", 12)
                    pdf.set_x(1.5)
                    pdf.multi_cell(
                        6.0, pdf.line_height, content, 
                        align="L", new_x="LMARGIN", new_y="NEXT"
                    )
                    pdf.ln(pdf.line_height)
        pdf.output(output_path)

    def _add_page_number(self, run):
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

    def render_docx(
        self, scenes: list[dict[str, Any]], output_path: str, 
        pre_scene_text: str = "", project_title: str = "Untitled"
    ):
        doc = Document()
        
        # Set Global Style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Courier'
        font.size = Pt(12)
        
        section = doc.sections[0]
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.header_distance = Inches(0.5)

        # Enable different header for the first page (title page)
        section.different_first_page_header_footer = True
        
        # Add page numbers to header (except first page)
        header = section.header
        p_header = header.paragraphs[0]
        p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_header = p_header.add_run()
        run_header.font.name = 'Courier'
        run_header.font.size = Pt(12)
        self._add_page_number(run_header)
        run_header.add_text(".")

        title_lines, teaser_lines = self._split_pre_scene(pre_scene_text)

        # 1. Title Page
        if not title_lines:
            title_lines = [project_title]

        # Push down approx 1/3
        for _ in range(15):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)

        for i, line in enumerate(title_lines):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().upper() if i == 0 else line.strip())
            run.font.name = 'Courier'
            if i == 0:
                run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
        
        doc.add_page_break()

        # 2. Script Start
        if teaser_lines:
            for line in teaser_lines:
                p = doc.add_paragraph()
                if line.strip():
                    run = p.add_run(self.sanitize(line))
                    run.font.name = 'Courier'
                p.paragraph_format.space_after = Pt(0)
            doc.add_paragraph()

        for scene in scenes:
            for elem in scene.get("elements", []):
                etype = elem.get("element_type", "action")
                content = self.sanitize(elem.get("content", ""))
                p = doc.add_paragraph()
                if etype == "scene_heading":
                    run = p.add_run(content.upper())
                    run.bold = True
                    run.font.name = 'Courier'
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(12)
                elif etype == "character":
                    run = p.add_run(content.upper())
                    run.font.name = 'Courier'
                    p.paragraph_format.left_indent = Inches(2.0)
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(0)
                elif etype == "parenthetical":
                    run = p.add_run(f"({content})")
                    run.font.name = 'Courier'
                    p.paragraph_format.left_indent = Inches(1.5)
                    p.paragraph_format.space_after = Pt(0)
                elif etype == "dialogue":
                    run = p.add_run(content)
                    run.font.name = 'Courier'
                    p.paragraph_format.left_indent = Inches(1.0)
                    p.paragraph_format.right_indent = Inches(1.5)
                    p.paragraph_format.space_after = Pt(12)
                elif etype == "transition":
                    run = p.add_run(content.upper())
                    run.font.name = 'Courier'
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(12)
                else: # action
                    run = p.add_run(content)
                    run.font.name = 'Courier'
                    p.paragraph_format.space_after = Pt(12)
        doc.save(output_path)
