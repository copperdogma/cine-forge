from __future__ import annotations

import re
from pathlib import Path

import pytest

from cine_forge.modules.ingest.story_ingest_v1.main import (
    _extract_pdf_text,
    classify_format,
    classify_format_with_diagnostics,
    detect_file_format,
    read_source_text,
    read_source_text_with_diagnostics,
    run_module,
)


@pytest.mark.unit
def test_detect_file_format_supports_expected_extensions() -> None:
    assert detect_file_format(Path("a.txt")) == "txt"
    assert detect_file_format(Path("a.md")) == "md"
    assert detect_file_format(Path("a.fountain")) == "fountain"
    assert detect_file_format(Path("a.fdx")) == "fdx"
    assert detect_file_format(Path("a.pdf")) == "pdf"
    assert detect_file_format(Path("a.docx")) == "docx"


@pytest.mark.unit
def test_detect_file_format_rejects_unknown_extension() -> None:
    with pytest.raises(ValueError, match="Unsupported input format"):
        detect_file_format(Path("a.xlsx"))


@pytest.mark.unit
def test_read_source_text_preserves_plain_text_exactly(tmp_path: Path) -> None:
    source = tmp_path / "sample.txt"
    payload = "Line one\nLine two\n"
    source.write_text(payload, encoding="utf-8")

    assert read_source_text(source) == payload


@pytest.mark.unit
def test_read_source_text_extracts_pdf_via_reader(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    source = tmp_path / "sample.pdf"
    source.write_bytes(b"%PDF-1.4 mock")

    monkeypatch.setattr(
        "cine_forge.modules.ingest.story_ingest_v1.main._extract_pdf_text_with_diagnostics",
        lambda _: ("INT. ROOM - NIGHT\nMARA\nHello there.", {}),
    )

    assert read_source_text(source) == "INT. ROOM - NIGHT\nMARA\nHello there."


@pytest.mark.unit
def test_extract_pdf_text_prefers_layout_mode(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    source = tmp_path / "layout.pdf"
    source.write_bytes(b"%PDF-1.4 mock")

    class _FakePage:
        def extract_text(self, layout: bool = False) -> str:
            if layout:
                return "INT. ROOM - NIGHT\nMARA\nHello there."
            return "INT.ROOM-NIGHT"

    class _FakePDF:
        def __init__(self) -> None:
            self.pages = [_FakePage()]

        def __enter__(self) -> _FakePDF:
            return self

        def __exit__(self, *args: object) -> None:
            pass

    class _FakePlumber:
        def open(self, _path: str | Path) -> _FakePDF:
            return _FakePDF()

    monkeypatch.setattr("pdfplumber.open", _FakePlumber().open)

    extracted = _extract_pdf_text(source)
    assert extracted.startswith("INT. ROOM - NIGHT")


@pytest.mark.unit
def test_extract_pdf_text_uses_ocr_when_pdfplumber_is_sparse(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    source = tmp_path / "ocr-source.pdf"
    source.write_bytes(b"%PDF-1.4 mock")

    def _fake_which(name: str) -> str | None:
        if name == "ocrmypdf":
            return "/usr/local/bin/ocrmypdf"
        return None

    monkeypatch.setattr("cine_forge.modules.ingest.story_ingest_v1.main.shutil.which", _fake_which)

    class _FakePage:
        def extract_text(self, layout: bool = False) -> str:
            del layout
            return "y"  # Too sparse

    class _FakePDF:
        def __init__(self) -> None:
            self.pages = [_FakePage()]

        def __enter__(self) -> _FakePDF:
            return self

        def __exit__(self, *args: object) -> None:
            pass

    monkeypatch.setattr("pdfplumber.open", lambda _: _FakePDF())

    calls: list[list[str]] = []

    class _Result:
        def __init__(self, code: int, out: str = "") -> None:
            self.returncode = code
            self.stdout = out

    def _fake_run(args: list[str], capture_output: bool, text: bool, check: bool) -> _Result:
        del capture_output, text, check
        calls.append(args)
        exe = Path(args[0]).name
        if exe == "ocrmypdf":
            out_pdf = Path(args[-1])
            out_pdf.write_bytes(b"%PDF-1.4 ocr")
            return _Result(0, "")
        raise AssertionError(f"unexpected command: {args}")

    monkeypatch.setattr("cine_forge.modules.ingest.story_ingest_v1.main.subprocess.run", _fake_run)

    # Use a side-effect to mock the second pdfplumber call after OCR
    original_open = _FakePDF
    call_count = 0

    def _mock_open(_path: str | Path) -> _FakePDF:
        nonlocal call_count
        call_count += 1
        if call_count > 1:
            # Second call (after OCR) returns meaningful text
            p = _FakePage()
            p.extract_text = lambda layout=False: (
                "INT. ROOM - NIGHT\nMARA\n"
                "OCR recovered screenplay dialogue with enough readable spacing.\n"
                "Additional words keep extraction above the meaningfulness threshold.\n"
            )
            pdf = original_open()
            pdf.pages = [p]
            return pdf
        return original_open()

    monkeypatch.setattr("pdfplumber.open", _mock_open)

    extracted = _extract_pdf_text(source)
    assert "OCR recovered screenplay dialogue" in extracted
    assert any(Path(cmd[0]).name == "ocrmypdf" for cmd in calls)


@pytest.mark.unit
def test_read_source_text_pdf_reports_extractor_path_diagnostics(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    source = tmp_path / "diag.pdf"
    source.write_bytes(b"%PDF-1.4 mock")

    monkeypatch.setattr(
        "cine_forge.modules.ingest.story_ingest_v1.main._extract_pdf_text_via_pdfplumber",
        lambda _: "x",
    )
    monkeypatch.setattr(
        "cine_forge.modules.ingest.story_ingest_v1.main._extract_pdf_text_via_ocr",
        lambda _: "INT. ROOM - NIGHT\nMARA\nOCR path chosen with enough words to pass checks.",
    )

    def _fake_meaningful(text: str) -> bool:
        return "OCR path chosen" in text

    monkeypatch.setattr(
        "cine_forge.modules.ingest.story_ingest_v1.main._is_meaningful_pdf_text",
        _fake_meaningful,
    )

    _content, diagnostics = read_source_text_with_diagnostics(source)
    assert diagnostics["pdf_extractor_selected"] == "ocrmypdf"
    assert diagnostics["pdf_extractors_attempted"] == ["pdfplumber", "ocrmypdf"]
    assert diagnostics["pdf_extractor_output_lengths"]["pdfplumber"] == 1
    assert diagnostics["pdf_extractor_output_lengths"]["ocrmypdf"] > 20


@pytest.mark.unit
def test_read_source_text_extracts_docx_via_python_docx(tmp_path: Path) -> None:
    from docx import Document
    source = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("INT. LAB - NIGHT")
    doc.add_paragraph("MARA")
    doc.add_paragraph("Testing DOCX.")
    doc.save(str(source))

    # read_source_text uses read_source_text_with_diagnostics
    content, diagnostics = read_source_text_with_diagnostics(source)
    assert "INT. LAB - NIGHT" in content
    assert "MARA" in content
    assert "Testing DOCX." in content
    assert diagnostics["docx_extracted"] is True


@pytest.mark.unit
@pytest.mark.parametrize(
    ("fixture_name", "expected_format", "expected_marker"),
    [
        ("owl_creek_bridge.txt", "txt", "An Occurrence at Owl Creek Bridge"),
        ("owl_creek_bridge_excerpt.md", "md", "An Occurrence at Owl Creek Bridge"),
        ("run_like_hell_teaser.fountain", "fountain", "RUN LIKE HELL"),
        ("sample_script.fdx", "fdx", "FinalDraft"),
        ("sample_script.docx", "docx", "INT. LAB - NIGHT"),
        ("pit_and_pendulum.pdf", "pdf", "THE PIT"),
        ("patent_registering_votes_us272011_scan_5p.pdf", "pdf", "APPARATUS FOR REGISTERING"),
        ("run_like_hell_teaser_scanned_5p.pdf", "pdf", "RUN"),
    ],
)
def test_read_source_text_fixture_matrix_has_sane_extraction(
    fixture_name: str, expected_format: str, expected_marker: str,
) -> None:
    workspace_root = Path(__file__).resolve().parents[2]
    fixture_path = workspace_root / "tests" / "fixtures" / "ingest_inputs" / fixture_name
    content, diagnostics = read_source_text_with_diagnostics(fixture_path)

    assert detect_file_format(fixture_path) == expected_format
    assert expected_marker in content
    assert len(content.strip()) > 40
    assert "\x00" not in content

    # Basic semantic quality gate: extraction must preserve normal word boundaries.
    word_gap_count = len(re.findall(r"\b[A-Za-z]{3,}\s+[A-Za-z]{3,}\b", content))
    min_word_gaps = 3 if expected_format == "docx" else 5
    assert word_gap_count >= min_word_gaps

    if expected_format == "pdf":
        assert "reflow_applied" in diagnostics
        assert diagnostics["repaired_character_count"] >= 0
        assert diagnostics["original_character_count"] >= 0


@pytest.mark.unit
def test_classify_screenplay_detects_structural_signals() -> None:
    content = "\n".join(
        [
            "INT. CONTROL ROOM - NIGHT",
            "",
            "MARA",
            "(whispering)",
            "We still have one chance.",
            "CUT TO:",
            "EXT. RIDGE - NIGHT",
        ]
    )
    result = classify_format(content=content, file_format="fountain")
    assert result["detected_format"] == "screenplay"
    assert result["confidence"] >= 0.6
    assert result["evidence"]


@pytest.mark.unit
def test_classify_fdx_defaults_to_screenplay() -> None:
    xml_content = (
        "<FinalDraft><Content><Paragraph Type='Scene Heading'>"
        "<Text>INT. LAB - NIGHT</Text></Paragraph></Content></FinalDraft>"
    )
    result = classify_format(
        content=xml_content,
        file_format="fdx",
    )
    assert result["detected_format"] == "screenplay"
    assert result["confidence"] >= 0.9


@pytest.mark.unit
def test_classify_prose_detects_narrative_paragraphs() -> None:
    content = (
        "She crossed the market square before sunrise and watched the mist pull away from "
        "the river. The letter in her pocket felt heavier with each step, as if it could "
        "choose her future for her."
    )
    result = classify_format(content=content, file_format="txt")
    assert result["detected_format"] == "prose"
    assert result["confidence"] >= 0.3


@pytest.mark.unit
def test_classify_notes_detects_list_patterns() -> None:
    content = "\n".join(
        [
            "- Hero: Mara",
            "- Objective: restore comms",
            "1. Enter tower",
            "2. Trigger beacon",
            "3. Escape storm",
        ]
    )
    result = classify_format(content=content, file_format="md")
    assert result["detected_format"] == "notes"
    assert result["confidence"] >= 0.3


@pytest.mark.unit
def test_classification_confidence_higher_for_clear_signals() -> None:
    strong_screenplay = "\n".join(
        [
            "INT. ROOM - NIGHT",
            "MARA",
            "(quietly)",
            "We have to move.",
            "CUT TO:",
            "EXT. HILL - NIGHT",
        ]
    )
    ambiguous = "Fragments and ideas without clear scene structure or full narrative lines"

    strong = classify_format(content=strong_screenplay, file_format="fountain")
    weak = classify_format(content=ambiguous, file_format="txt")
    assert strong["confidence"] > weak["confidence"]


@pytest.mark.unit
def test_classify_tokenized_pdf_screenplay_detects_screenplay_signals() -> None:
    content = "\n".join(
        [
            "FADE",
            "IN:",
            "EXT.",
            "CITY",
            "CENTRE",
            "-",
            "NIGHT",
            "A",
            "ruined",
            "city.",
            "EXT.",
            "RUDDY",
            "&",
            "GREENE",
            "BUILDING",
            "-",
            "FRONT",
            "-",
            "NIGHT",
            "THE",
            "MARINER",
            "enters.",
        ]
        * 4
    )
    result, diagnostics = classify_format_with_diagnostics(content=content, file_format="pdf")
    assert result["detected_format"] == "screenplay"
    assert diagnostics["signals"]["tokenized_heading_sequences"] >= 2
    assert diagnostics["score_breakdown"]["screenplay"] > diagnostics["score_breakdown"]["prose"]


@pytest.mark.unit
def test_read_source_text_repairs_tokenized_pdf_layout(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    source = tmp_path / "tokenized.pdf"
    source.write_bytes(b"%PDF-1.4 mock")
    tokenized = "\n".join(
        [
            "FADE",
            "IN:",
            "EXT.",
            "CITY",
            "CENTRE",
            "-",
            "NIGHT",
            "A",
            "ruined",
            "city.",
            "CUT",
            "TO:",
            "EXT.",
            "DOCKS",
            "-",
            "DAWN",
            "The",
            "Mariner",
            "waits.",
        ]
        * 6
    )
    monkeypatch.setattr(
        "cine_forge.modules.ingest.story_ingest_v1.main._extract_pdf_text_with_diagnostics",
        lambda _: (tokenized, {}),
    )

    repaired, diagnostics = read_source_text_with_diagnostics(source)
    assert diagnostics["tokenized_layout_detected"] is True
    assert diagnostics["reflow_applied"] is True
    assert "EXT. CITY CENTRE - NIGHT" in repaired
    assert "CUT TO:" in repaired


@pytest.mark.unit
def test_read_source_text_repairs_compact_pdf_scene_headings(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    source = tmp_path / "compact.pdf"
    source.write_bytes(b"%PDF-1.4 mock")
    monkeypatch.setattr(
        "cine_forge.modules.ingest.story_ingest_v1.main._extract_pdf_text_with_diagnostics",
        lambda _: (
            "\n".join(
                [
                    "EXT.CITYCENTRE- NIGHT",
                    "Action line.",
                    "INT.RUDDY& GREENBUILDING- ELEVATOR",
                    "Another line.",
                    "BEGINFLASHBACK:EXT.COASTLINE- DAY- PAST",
                ]
            ),
            {},
        ),
    )

    repaired, diagnostics = read_source_text_with_diagnostics(source)
    assert "EXT. CITYCENTRE - NIGHT" in repaired
    assert "INT. RUDDY& GREENBUILDING - ELEVATOR" in repaired
    assert "BEGINFLASHBACK:\nEXT. COASTLINE - DAY - PAST" in repaired
    assert diagnostics["compact_heading_repairs"] >= 2


@pytest.mark.unit
def test_classify_compact_pdf_screenplay_detects_screenplay_signals() -> None:
    content = "\n".join(
        [
            "EXT. CITYCENTRE - NIGHT",
            "A ruined city.",
            "INT. RUDDY& GREENBUILDING - ELEVATOR",
            "ROSE",
            "Where are we?",
            "CUT TO:",
            "EXT. COASTLINE - DAY",
        ]
    )
    result = classify_format(content=content, file_format="pdf")
    assert result["detected_format"] in {"screenplay", "hybrid"}
    assert result["confidence"] >= 0.3


@pytest.mark.unit
def test_run_module_builds_raw_input_payload(tmp_path: Path) -> None:
    source = tmp_path / "notes.md"
    source_text = "- beat one\n- beat two\n"
    source.write_text(source_text, encoding="utf-8")

    result = run_module(
        inputs={},
        params={"input_file": str(source)},
        context={"run_id": "unit", "stage_id": "ingest", "runtime_params": {}},
    )

    artifact = result["artifacts"][0]
    assert artifact["artifact_type"] == "raw_input"
    assert artifact["data"]["content"] == source_text
    assert artifact["data"]["source_info"]["original_filename"] == "notes.md"
    assert artifact["data"]["classification"]["detected_format"] in {
        "notes",
        "unknown",
        "hybrid",
        "prose",
        "screenplay",
    }


@pytest.mark.unit
def test_run_module_records_classification_and_extraction_diagnostics(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    source = tmp_path / "tokenized.pdf"
    source.write_bytes(b"%PDF-1.4 mock")
    monkeypatch.setattr(
        "cine_forge.modules.ingest.story_ingest_v1.main._extract_pdf_text_with_diagnostics",
        lambda _: (
            "\n".join(
                [
                    "EXT.",
                    "CITY",
                    "CENTRE",
                    "-",
                    "NIGHT",
                    "THE",
                    "MARINER",
                    "moves",
                    "EXT.",
                    "DOCKS",
                    "-",
                    "DAWN",
                    "He",
                    "stops",
                ]
                * 8
            ),
            {},
        ),
    )

    result = run_module(
        inputs={},
        params={"input_file": str(source)},
        context={"run_id": "unit", "stage_id": "ingest", "runtime_params": {}},
    )
    metadata = result["artifacts"][0]["metadata"]
    diagnostics = metadata["annotations"]
    assert "classification_diagnostics" in diagnostics
    assert diagnostics["classification_diagnostics"]["signals"]["scene_headings"] >= 2
    assert result["artifacts"][0]["data"]["classification"]["detected_format"] != "unknown"
    assert diagnostics["extraction_diagnostics"]["tokenized_layout_detected"] is True


@pytest.mark.unit
def test_run_module_rejects_empty_extracted_source_text(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path,
) -> None:
    source = tmp_path / "empty.pdf"
    source.write_bytes(b"%PDF-1.4 mock")
    monkeypatch.setattr(
        "cine_forge.modules.ingest.story_ingest_v1.main.read_source_text_with_diagnostics",
        lambda _: ("", {"pdf_extractor_selected": "fallback_sparse"}),
    )

    with pytest.raises(ValueError, match="could not extract readable text"):
        run_module(
            inputs={},
            params={"input_file": str(source)},
            context={"run_id": "unit", "stage_id": "ingest", "runtime_params": {}},
        )
